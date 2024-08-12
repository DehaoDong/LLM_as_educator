from docx import Document
import pandas as pd
import json
import prompt_engineering as pe

# Step 1: Load the rubric text
rubric_doc = Document('dataset/ethics-assignment-2024.docx')
rubric_text = ""
for paragraph in rubric_doc.paragraphs:
    rubric_text += paragraph.text + "\n"

# Step 2: Load assignments for students 1 to 16
assignments = {}
for i in range(1, 17):
    file_path = f'dataset/data/student {i}.docx'
    doc = Document(file_path)
    assignment_text = ""
    for paragraph in doc.paragraphs:
        assignment_text += paragraph.text + "\n"
    assignments[f'student_{i}'] = assignment_text

# Step 3: Load markings from the Excel file
df = pd.read_excel('dataset/class-list-marking.xlsx')
selected_columns = df[['anon', 'Feedback for BB', 'Total/100', 'description/20',
                       'roles/10', 'risks/40', 'mitigation/20', 'refs/10']]
first_30_students = selected_columns.head(30)

# Step 4: Create the dataset with prompts
dataset = []

# 16 files
for i in range(1, 17):
    # Get the assignment text for the current student
    assignment_text = assignments.get(f'student_{i}', '')

    # Fill the rubric and assignment into the prompt template
    filled_prompt = pe.FE_PROMPT_TEMPLATE.format(rubric=rubric_text, assignment=assignment_text)

    # Separate the prompt into system and user prompts
    system_prompt = filled_prompt.split("<</SYS>>")[0].replace("<<SYS>>", "").strip()
    user_prompt = filled_prompt.split("<<USR>>")[1].replace("<</USR>>", "").strip()

    # Retrieve the marking for the current student
    student_marking = first_30_students.iloc[i-1]
    assistant_response = {
        "feedback": student_marking['Feedback for BB'],
        "description / 20": int(student_marking['description/20']),
        "roles / 10": int(student_marking['roles/10']),
        "risks / 40": int(student_marking['risks/40']),
        "mitigation / 20": int(student_marking['mitigation/20']),
        "refs / 10": int(student_marking['refs/10']),
        "total score": float(student_marking['Total/100'])  # Convert to float if needed
    }

    # Create a conversation dataset structure
    conversation = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
        {"role": "assistant", "content": json.dumps(assistant_response, indent=4)}
    ]

    # Append the conversation to the dataset
    dataset.append(conversation)

# Convert the dataset to JSON format
json_dataset = json.dumps(dataset, indent=4)

# Save the JSON dataset to a file
with open('../generated_dataset.json', 'w') as f:
    f.write(json_dataset)

# Optionally print the first conversation in the dataset
# print(json_dataset[:500])  # Truncated output for brevity
